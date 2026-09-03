from pathlib import Path
from docx import Document


SOURCE = Path(r"C:\Users\0009\.gemini\antigravity\scratch\schedule-viewer\.montage_work\Повний_курс_лекцій_Монтаж_та_обслуговування_КС_2025_2026.docx")

REPLACEMENTS = {
    "TIA/EIA-568-B": "ANSI/TIA-568.2-E (схема T568B)",
    "NIST SP 800-88 Rev.1": "NIST SP 800-88 Rev. 2",
    "NIST Special Publication 800-88 Revision 1": "NIST Special Publication 800-88 Revision 2",
    "National Institute of Standards and Technology, U.S. Department of Commerce, 2020.":
        "National Institute of Standards and Technology, 2025.",
}


def iter_paragraphs(document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def replace_in_paragraph(paragraph, old, new):
    if old not in paragraph.text:
        return 0
    full = "".join(run.text for run in paragraph.runs)
    if old not in full:
        raise RuntimeError(f"Text is split outside ordinary runs: {old}")
    updated = full.replace(old, new)
    paragraph.runs[0].text = updated
    for run in paragraph.runs[1:]:
        run.text = ""
    return full.count(old)


doc = Document(SOURCE)
counts = {key: 0 for key in REPLACEMENTS}
for paragraph in iter_paragraphs(doc):
    for old, new in REPLACEMENTS.items():
        counts[old] += replace_in_paragraph(paragraph, old, new)

missing = [old for old, count in counts.items() if count == 0]
if missing:
    raise RuntimeError(f"Expected text not found: {missing}")

doc.save(SOURCE)
print(counts)
